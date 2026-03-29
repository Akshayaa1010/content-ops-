import uuid, os
from pypdf import PdfReader
from app.db.database import AsyncSessionLocal
from app.db.models import Document, DocumentChunk
from app.pipeline.embeddings import embed_texts
from sqlalchemy import select
from app.agents.llm_client import generate_json
import json

async def process_document(file_path: str, document_id: str):
    full_text = ""
    ext = file_path.split(".")[-1].lower()
    
    if ext == "pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            full_text += page.extract_text() + "\n"
    elif ext == "docx":
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            full_text = "ERROR: python-docx not installed"
    else: # Default to plain text
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()

    # Minimal chunking logic: 1000 chars with 200 overlap
    chunk_size = 1000
    overlap = 200
    chunks = []
    for i in range(0, len(full_text), chunk_size - overlap):
        chunks.append(full_text[i:i + chunk_size])

    async with AsyncSessionLocal() as db:
        # Fetch existing document
        result = await db.execute(select(Document).where(Document.id == uuid.UUID(document_id)))
        doc = result.scalar_one_or_none()
        if not doc:
            raise ValueError(f"Document {document_id} not found")

        doc.raw_text = full_text
        doc.chunk_count = len(chunks)
        await db.flush()

        # Embed chunks in batches
        embeddings = await embed_texts(chunks)

        for i, (text, emb) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                id=uuid.uuid4(),
                document_id=doc.id,
                chunk_text=text,
                embedding=emb,
                chunk_index=i
            )
            db.add(chunk)

        await db.commit()
        return doc.id

async def index_policy_file(file_path: str):
    if not os.path.exists(file_path):
        return
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Split by numbered rules (e.g. "1. Tone", "2. Claims")
    import re
    rules = re.split(r'\n(?=\d+\.)', content)
    
    async with AsyncSessionLocal() as db:
        # Create a special "system" document for policies
        doc_id = uuid.uuid5(uuid.NAMESPACE_DNS, "brand_policies")
        result = await db.execute(select(Document).where(Document.id == doc_id))
        doc = result.scalar_one_or_none()
        if not doc:
            doc = Document(id=doc_id, filename="brand_policies.txt", doc_type="policy")
            db.add(doc)
            await db.flush()
        
        doc.raw_text = content
        doc.chunk_count = len(rules)
        
        # Clear old chunks
        from app.db.models import DocumentChunk
        from sqlalchemy import delete
        await db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == doc_id))
        
        embeddings = await embed_texts(rules)
        for i, (text, emb) in enumerate(zip(rules, embeddings)):
            chunk = DocumentChunk(
                document_id=doc_id,
                chunk_text=text,
                embedding=emb,
                chunk_index=i
            )
            db.add(chunk)
        await db.commit()

async def search_knowledge(query_embedding: list[float], limit: int = 5):
    from pgvector.sqlalchemy import Vector
    async with AsyncSessionLocal() as db:
        # Simple vector similarity search
        result = await db.execute(
            select(DocumentChunk)
            .order_by(DocumentChunk.embedding.l2_distance(query_embedding))
            .limit(limit)
        )
        return [row.chunk_text for row in result.scalars().all()]

async def generate_content(text: str) -> dict:
    system_prompt = """
    You are a world-class marketing expert.
    Convert the provided product specification into a marketing blog and a LinkedIn post.
    
    Output MUST be a valid JSON object with the following keys:
    - "blog": A detailed marketing blog.
    - "linkedin_post": A concise, engaging LinkedIn post with hashtags.
    """
    
    user_message = f"Convert the following product specification into a marketing blog and a LinkedIn post:\n\n{text}"
    response_json = await generate_json(system_prompt, user_message)
    return json.loads(response_json)
